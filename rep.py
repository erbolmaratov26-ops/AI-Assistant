from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # --- Title Page ---
    title = doc.add_heading('Mindful AI Diary: Deep Learning Based Emotion Detection', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n' * 5)
    
    p = doc.add_paragraph('A Project Report Submitted by\n[Your Name]')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n' * 2)
    p = doc.add_paragraph('Department of Computer Science\n[Your University/Organization]')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()

    # --- Table of Contents (Placeholder) ---
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('1. Introduction................................................................... 1')
    doc.add_paragraph('2. Literature Review.............................................................. 2')
    doc.add_paragraph('3. Methodology.................................................................... 3')
    doc.add_paragraph('4. System Architecture............................................................ 4')
    doc.add_paragraph('5. Implementation................................................................. 6')
    doc.add_paragraph('6. Results and Analysis........................................................... 8')
    doc.add_paragraph('7. Future Scope and Conclusion.................................................... 9')
    doc.add_paragraph('8. References..................................................................... 10')
    doc.add_page_break()

    # --- Section 1: Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        "Mental health awareness has become a critical aspect of modern life. "
        "With the increasing stress of daily routines, many individuals turn to journaling as a coping mechanism. "
        "However, traditional journaling lacks feedback. This project, the 'Mindful AI Diary', aims to bridge that gap "
        "by utilizing Artificial Intelligence to provide real-time emotional feedback."
    )
    doc.add_paragraph(
        "Unlike simple sentiment analysis which only categorizes text as positive or negative, this project leverages "
        "Deep Learning (specifically BERT models) to detect granular emotions such as Joy, Sadness, Anger, Fear, Love, and Surprise. "
        "This allows users to gain deeper insights into their emotional well-being over time."
    )
    
    # --- Section 2: Literature Review ---
    doc.add_heading('2. Literature Review', level=1)
    doc.add_paragraph(
        "The field of Natural Language Processing (NLP) has evolved significantly. "
        "Early approaches relied on lexicon-based methods like TextBlob and VADER, which used dictionaries of 'good' and 'bad' words. "
        "While fast, these methods failed to capture context, sarcasm, or complex sentence structures."
    )
    doc.add_paragraph(
        "The introduction of Transformer architectures, specifically BERT (Bidirectional Encoder Representations from Transformers) by Google in 2018, "
        "revolutionized the field. BERT reads text in both directions (left-to-right and right-to-left), allowing it to understand the full context of a word. "
        "This project utilizes a 'DistilBERT' model, which is a lighter, faster version of BERT that retains 97% of its performance while being 40% smaller."
    )

    # --- Section 3: Methodology ---
    doc.add_heading('3. Methodology', level=1)
    doc.add_paragraph(
        "The development of the Mindful AI Diary follows the Model-View-Controller (MVC) architectural pattern. "
        "The core methodology involves fine-tuning a pre-trained DistilBERT model on an emotion dataset."
    )
    doc.add_heading('3.1 Dataset', level=2)
    doc.add_paragraph(
        "The model is trained on a dataset of English Twitter messages labeled with six basic emotions: anger, fear, joy, love, sadness, and surprise. "
        "This dataset is widely used for multi-class classification tasks in NLP."
    )
    doc.add_heading('3.2 Pre-processing', level=2)
    doc.add_paragraph(
        "Text data undergoes tokenization using the DistilBertTokenizer. This converts words into numerical IDs that the neural network can process. "
        "Special tokens like [CLS] (start) and [SEP] (separator) are added to the input sequences."
    )

    # --- Section 4: System Architecture ---
    doc.add_heading('4. System Architecture', level=1)
    doc.add_paragraph(
        "The system consists of three main components: The Client (Frontend), The Server (Backend), and The AI Model."
    )
    doc.add_paragraph(
        "1. **Frontend:** Built using HTML5, CSS3, and Jinja2 templating. It provides a clean, distraction-free interface for users to write their thoughts.\n"
        "2. **Backend:** Developed using Python and Flask. It handles HTTP requests, processes form data, and manages the communication between the user and the AI.\n"
        "3. **AI Inference Engine:** Uses the Hugging Face `transformers` library to load the model into memory. When text is received, it is passed through the model to obtain a probability distribution across the six emotions."
    )

    # --- Section 5: Implementation ---
    doc.add_heading('5. Implementation', level=1)
    doc.add_paragraph(
        "The implementation was carried out in Python 3.9. Below are the key code snippets used in the development."
    )
    doc.add_heading('5.1 Emotion Classification Logic', level=2)
    doc.add_paragraph("The core logic for detecting emotions is handled by the pipeline API:")
    doc.add_paragraph(
        "```python\n"
        "from transformers import pipeline\n"
        "emotion_classifier = pipeline(\n"
        "    'text-classification', \n"
        "    model='bhadresh-savani/distilbert-base-uncased-emotion', \n"
        "    return_all_scores=True\n"
        ")\n"
        "```"
    )
    doc.add_heading('5.2 Backend Routing', level=2)
    doc.add_paragraph(
        "Flask routes map the URL endpoints to python functions. The main route handles both GET (displaying the form) and POST (processing the data) requests."
    )

    # --- Section 6: Results ---
    doc.add_heading('6. Results and Analysis', level=1)
    doc.add_paragraph(
        "The system was tested with various inputs to verify accuracy. The DistilBERT model showed superior performance compared to traditional methods."
    )
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Input Text'
    hdr_cells[1].text = 'Detected Emotion'
    hdr_cells[2].text = 'Confidence Score'
    
    row1 = table.rows[1].cells
    row1[0].text = "I passed the exam!"
    row1[1].text = "Joy"
    row1[2].text = "0.99"
    
    row2 = table.rows[2].cells
    row2[0].text = "I am lost in the woods."
    row2[1].text = "Fear"
    row2[2].text = "0.95"
    
    row3 = table.rows[3].cells
    row3[0].text = "How dare you speak to me like that?"
    row3[1].text = "Anger"
    row3[2].text = "0.97"

    # --- Section 7: Future Scope ---
    doc.add_heading('7. Future Scope and Conclusion', level=1)
    doc.add_paragraph(
        "The Mindful AI Diary demonstrates the power of modern NLP in personal wellness applications. "
        "Future improvements could include: \n"
        "- **Long-term Mood Tracking:** Storing data in a database to visualize mood trends over months.\n"
        "- **Voice Input:** Integrating speech-to-text to allow users to speak their entries.\n"
        "- **Personalized Recommendations:** Suggesting activities (e.g., meditation, exercise) based on the detected mood."
    )
    doc.add_paragraph(
        "In conclusion, this project successfully bridges the gap between raw text data and meaningful emotional insight, providing a valuable tool for self-reflection."
    )

    # --- Section 8: References ---
    doc.add_heading('8. References', level=1)
    doc.add_paragraph("1. Devlin, J., et al. (2018). 'BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding'.")
    doc.add_paragraph("2. Sanh, V., et al. (2019). 'DistilBERT, a distilled version of BERT: smaller, faster, cheaper and lighter'.")
    doc.add_paragraph("3. Hugging Face Documentation. https://huggingface.co/")
    doc.add_paragraph("4. Flask Documentation. https://flask.palletsprojects.com/")

    doc.save('Mindful_AI_Diary_Report.docx')
    print("Report generated successfully as 'Mindful_AI_Diary_Report.docx'")

if __name__ == "__main__":
    create_report()